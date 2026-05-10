"""
base.py

Общий модуль для raw-парсеров облачных провайдеров.

Здесь лежит всё общее:
- загрузка HTML-страниц;
- извлечение текста и ссылок;
- скачивание PDF/DOCX/XLSX/CSV;
- извлечение текста из документов;
- сохранение JSON;
- базовые dataclass-структуры;
- parse_log.

Этот файл НЕ содержит логику конкретного провайдера.
Логика T1, Selectel, VK Cloud, Cloud.ru должна лежать в отдельных файлах.
"""

from __future__ import annotations

import json
import logging
import re
import time
from dataclasses import asdict, dataclass, field, is_dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup


# Чтобы pypdf не засорял консоль предупреждениями вида:
# Ignoring wrong pointing object 6 0 (offset 0)
logging.getLogger("pypdf").setLevel(logging.ERROR)


REQUEST_TIMEOUT = 30
REQUEST_DELAY_SECONDS = 1.0

DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 "
        "(compatible; CloudMarketplaceStudentParser/1.0; "
        "educational project; public pages only)"
    )
}

DOCUMENT_EXTENSIONS = (
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".csv",
)


@dataclass
class ParseLogItem:
    provider_id: str
    url: str
    parsed_at: str
    status: str
    records_added: int
    error: str | None = None


@dataclass
class RawSource:
    provider_id: str
    source_type: str
    url: str
    final_url: str | None
    http_status: int | None
    title: str | None
    headings: list[str]
    text: str
    links: list[dict[str, str]]
    document_links: list[dict[str, str]]
    downloaded_documents: list[dict[str, Any]]
    parsed_at: str


@dataclass
class RawDataset:
    provider: dict[str, Any]
    collected_at: str
    source_pages: list[RawSource] = field(default_factory=list)
    service_candidates: list[dict[str, Any]] = field(default_factory=list)
    pricing_sources: list[dict[str, Any]] = field(default_factory=list)
    docs_sources: list[dict[str, Any]] = field(default_factory=list)
    api_sources: list[dict[str, Any]] = field(default_factory=list)
    security_sources: list[dict[str, Any]] = field(default_factory=list)
    pricing_items_raw: list[dict[str, Any]] = field(default_factory=list)
    compliance_evidence_raw: list[dict[str, Any]] = field(default_factory=list)
    region_evidence_raw: list[dict[str, Any]] = field(default_factory=list)
    parse_log: list[ParseLogItem] = field(default_factory=list)

def now_iso() -> str:
    """
    Возвращает текущее время в ISO-формате с timezone UTC.

    Это значение нужно писать в parsed_at / collected_at.
    """
    return datetime.now(timezone.utc).isoformat()


def normalize_spaces(text: str | None) -> str:
    """
    Убирает лишние пробелы, переносы строк и табы.
    """
    if not text:
        return ""

    return re.sub(r"\s+", " ", text).strip()


def is_document_url(url: str) -> bool:
    """
    Проверяет, является ли ссылка документом: PDF, DOCX, XLSX, CSV.
    """
    path = urlparse(url).path.lower()
    return path.endswith(DOCUMENT_EXTENSIONS)


def is_internal_url(url: str, allowed_domains: set[str]) -> bool:
    """
    Проверяет, относится ли ссылка к нужному домену.

    allowed_domains пример:
        {"t1-cloud.ru", "www.t1-cloud.ru"}

    Пустой netloc разрешаем, потому что относительные ссылки тоже внутренние.
    """
    parsed = urlparse(url)
    return parsed.netloc in allowed_domains or parsed.netloc == ""


def safe_filename(url: str) -> str:
    """
    Создаёт безопасное имя файла из URL.
    """
    parsed = urlparse(url)
    name = Path(parsed.path).name or "document"

    name = re.sub(r"[^a-zA-Zа-яА-Я0-9_.-]+", "_", name)

    return name[:180]


def fetch(
    url: str,
    headers: dict[str, str] | None = None,
    timeout: int = REQUEST_TIMEOUT,
    delay_seconds: float = REQUEST_DELAY_SECONDS,
) -> requests.Response:
    """
    Загружает страницу или файл.

    Важно:
    - делает паузу между запросами;
    - подставляет User-Agent;
    - исправляет кодировку HTML-страниц.
    """
    time.sleep(delay_seconds)

    response = requests.get(
        url,
        headers=headers or DEFAULT_HEADERS,
        timeout=timeout,
        allow_redirects=True,
    )

    response.raise_for_status()

    content_type = response.headers.get("Content-Type", "").lower()

    if "text/html" in content_type:
        response.encoding = response.apparent_encoding or "utf-8"

    return response


def deduplicate_links(links: list[dict[str, str]]) -> list[dict[str, str]]:
    """
    Убирает дубликаты ссылок.
    """
    seen = set()
    result = []

    for link in links:
        url = link.get("url")

        if not url or url in seen:
            continue

        seen.add(url)
        result.append(link)

    return result


def extract_html(
    html: str,
    page_url: str,
    allowed_domains: set[str],
) -> tuple[str | None, list[str], str, list[dict[str, str]], list[dict[str, str]]]:
    """
    Извлекает из HTML:
    - title;
    - headings;
    - основной текст;
    - все внутренние ссылки;
    - ссылки на документы.

    Возвращает:
        title, headings, text, links, document_links
    """

    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()

    title = normalize_spaces(soup.title.get_text(" ")) if soup.title else None

    headings = []
    for tag in soup.find_all(["h1", "h2", "h3"]):
        heading_text = normalize_spaces(tag.get_text(" "))

        if heading_text:
            headings.append(heading_text)

    main = soup.find("main") or soup.find("article") or soup.body or soup
    page_text = normalize_spaces(main.get_text(" "))

    links = []
    document_links = []

    for a in soup.find_all("a", href=True):
        href = a.get("href")
        absolute_url = urljoin(page_url, href)
        anchor = normalize_spaces(a.get_text(" "))

        if not is_internal_url(absolute_url, allowed_domains):
            continue

        item = {
            "anchor": anchor,
            "url": absolute_url,
        }

        links.append(item)

        if is_document_url(absolute_url):
            document_links.append(item)

    return (
        title,
        headings,
        page_text,
        deduplicate_links(links),
        deduplicate_links(document_links),
    )


def extract_pdf_text(file_path: Path) -> str:
    """
    Извлекает текст из PDF.

    По умолчанию используем pypdf.
    Если PDF плохо читается, можно заменить на PyMuPDF.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    parts = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_spaces(page.extract_text())

        if text:
            parts.append(f"[PDF_PAGE {page_index}] {text}")

    return "\n".join(parts)


def extract_docx_text(file_path: Path) -> str:
    """
    Извлекает текст из DOCX.
    """
    from docx import Document

    document = Document(str(file_path))
    parts = []

    for paragraph in document.paragraphs:
        text = normalize_spaces(paragraph.text)

        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"[DOCX_TABLE {table_index}]")

        for row in table.rows:
            cells = [normalize_spaces(cell.text) for cell in row.cells]

            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_xlsx_text(file_path: Path) -> str:
    """
    Извлекает текст из XLSX.
    """
    from openpyxl import load_workbook

    workbook = load_workbook(str(file_path), data_only=True)
    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"[XLSX_SHEET] {sheet.title}")

        for row in sheet.iter_rows(values_only=True):
            values = [
                normalize_spaces(str(value))
                for value in row
                if value is not None
            ]

            if values:
                parts.append(" | ".join(values))

    return "\n".join(parts)


def extract_document_text(file_path: Path) -> str:
    """
    Определяет тип документа по расширению и извлекает текст.
    """
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            return extract_pdf_text(file_path)

        if suffix == ".docx":
            return extract_docx_text(file_path)

        if suffix == ".xlsx":
            return extract_xlsx_text(file_path)

        if suffix == ".csv":
            return file_path.read_text(encoding="utf-8", errors="ignore")

        return ""

    except Exception as exc:
        return f"[EXTRACTION_FAILED] {exc}"


def download_document(
    url: str,
    provider_id: str,
    output_dir: Path,
) -> dict[str, Any]:
    """
    Скачивает документ и извлекает из него текст.

    output_dir пример:
        Path("data/raw/downloads")

    Файл сохранится в:
        data/raw/downloads/<provider_id>/
    """
    provider_downloads_dir = output_dir / provider_id
    provider_downloads_dir.mkdir(parents=True, exist_ok=True)

    try:
        response = fetch(url)

        filename = safe_filename(response.url)
        content_type = response.headers.get("Content-Type", "")

        if "." not in filename:
            low_type = content_type.lower()

            if "pdf" in low_type:
                filename += ".pdf"
            elif "word" in low_type:
                filename += ".docx"
            elif "excel" in low_type or "spreadsheet" in low_type:
                filename += ".xlsx"
            else:
                filename += ".bin"

        file_path = provider_downloads_dir / filename
        file_path.write_bytes(response.content)

        extracted_text = extract_document_text(file_path)

        return {
            "url": url,
            "final_url": response.url,
            "http_status": response.status_code,
            "local_path": str(file_path),
            "content_type": content_type,
            "size_bytes": len(response.content),
            "extracted_text": extracted_text,
            "status": "success",
            "error": None,
        }

    except Exception as exc:
        return {
            "url": url,
            "final_url": None,
            "http_status": None,
            "local_path": None,
            "content_type": None,
            "size_bytes": None,
            "extracted_text": "",
            "status": "failed",
            "error": str(exc),
        }


def collect_html_page(
    provider_id: str,
    source_type: str,
    url: str,
    allowed_domains: set[str],
    downloads_output_dir: Path,
    download_documents: bool = True,
) -> tuple[RawSource | None, ParseLogItem]:
    """
    Собирает одну HTML-страницу.

    Возвращает:
        RawSource | None
        ParseLogItem

    Если страница успешно собрана, status будет success.
    Если нет — failed.
    """
    parsed_at = now_iso()

    try:
        response = fetch(url)

        title, headings, text, links, document_links = extract_html(
            html=response.text,
            page_url=response.url,
            allowed_domains=allowed_domains,
        )

        downloaded_documents = []

        if download_documents:
            for document_link in document_links:
                downloaded = download_document(
                    url=document_link["url"],
                    provider_id=provider_id,
                    output_dir=downloads_output_dir,
                )

                downloaded["anchor"] = document_link.get("anchor", "")
                downloaded_documents.append(downloaded)

        raw_source = RawSource(
            provider_id=provider_id,
            source_type=source_type,
            url=url,
            final_url=response.url,
            http_status=response.status_code,
            title=title,
            headings=headings,
            text=text,
            links=links,
            document_links=document_links,
            downloaded_documents=downloaded_documents,
            parsed_at=parsed_at,
        )

        log = ParseLogItem(
            provider_id=provider_id,
            url=url,
            parsed_at=parsed_at,
            status="success",
            records_added=1,
            error=None,
        )

        return raw_source, log

    except Exception as exc:
        log = ParseLogItem(
            provider_id=provider_id,
            url=url,
            parsed_at=parsed_at,
            status="failed",
            records_added=0,
            error=str(exc),
        )

        return None, log


def to_jsonable(data: Any) -> Any:
    """
    Преобразует dataclass-объекты в обычные dict/list для json.dump().
    """
    if is_dataclass(data):
        return asdict(data)

    if isinstance(data, list):
        return [to_jsonable(item) for item in data]

    if isinstance(data, dict):
        return {key: to_jsonable(value) for key, value in data.items()}

    return data


def save_json(path: Path, data: Any) -> None:
    """
    Сохраняет данные в JSON.
    """
    path.parent.mkdir(parents=True, exist_ok=True)

    with path.open("w", encoding="utf-8") as file:
        json.dump(
            to_jsonable(data),
            file,
            ensure_ascii=False,
            indent=2,
        )


def dataset_to_dict(dataset: RawDataset) -> dict[str, Any]:
    """
    Превращает RawDataset в словарь для сохранения.
    """
    return {
        "provider": dataset.provider,
        "collected_at": dataset.collected_at,
        "source_pages": [asdict(page) for page in dataset.source_pages],
        "service_candidates": dataset.service_candidates,
        "pricing_sources": dataset.pricing_sources,
        "docs_sources": dataset.docs_sources,
        "api_sources": dataset.api_sources,
        "security_sources": dataset.security_sources,
        "pricing_items_raw": dataset.pricing_items_raw,
        "compliance_evidence_raw": dataset.compliance_evidence_raw,
        "region_evidence_raw": dataset.region_evidence_raw,
        "parse_log": [asdict(item) for item in dataset.parse_log],
    }


def print_dataset_summary(dataset: RawDataset, raw_path: Path, log_path: Path) -> None:
    """
    Красивый вывод результата в консоль.
    """
    print("Done.")
    print(f"Raw data: {raw_path}")
    print(f"Parse log: {log_path}")
    print(f"Source pages: {len(dataset.source_pages)}")
    print(f"Service candidates: {len(dataset.service_candidates)}")
    print(f"Pricing sources: {len(dataset.pricing_sources)}")
    print(f"Docs sources: {len(dataset.docs_sources)}")
    print(f"API sources: {len(dataset.api_sources)}")
    print(f"Security sources: {len(dataset.security_sources)}")
    print(f"Pricing items raw: {len(dataset.pricing_items_raw)}")
    print(f"Compliance evidence raw: {len(dataset.compliance_evidence_raw)}")
    print(f"Region evidence raw: {len(dataset.region_evidence_raw)}")